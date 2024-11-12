import os
from openai import OpenAI
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException
from dotenv import load_dotenv
import time
from datetime import datetime
import docx

# Load environment variables from .env file
load_dotenv()

# Credentials
naver_marketing_id = os.getenv("NAVER_MARKETING_ID")
naver_marketing_password = os.getenv("NAVER_MARKETING_PASSWORD")
OpenAI.api_key = os.getenv("OPENAI_API_KEY")  # Ensure your OpenAI API key is in the .env file

# Set up WebDriver
options = Options()
options.add_argument("--disable-notifications")
service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service, options=options)
driver.maximize_window()

def login_to_naver():
    url = "https://nid.naver.com/nidlogin.login"
    driver.get(url)
    try:
        WebDriverWait(driver, 10).until(lambda d: d.find_element(By.ID, "id")).send_keys(naver_marketing_id)
        WebDriverWait(driver, 10).until(lambda d: d.find_element(By.ID, "pw")).send_keys(naver_marketing_password)
        WebDriverWait(driver, 10).until(lambda d: d.find_element(By.ID, "log.login")).click()
        time.sleep(20)  # Allow time for CAPTCHA
    except TimeoutException:
        print("Login elements did not load in time.")

def scrape_posts(search_keyword, start_page=1, end_page=3):
    scraped_data = []
    filter_date = datetime(2024, 11, 1)  # Set the filter date to November 1, 2024

    try:
        search_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//input[@name="query"]'))
        )
        search_input.clear()
        search_input.send_keys(search_keyword)

        search_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//button[@onclick="searchBoard();return false;"]'))
        )
        search_button.click()
        time.sleep(2)

        # Switch to the required iframe
        WebDriverWait(driver, 10).until(EC.frame_to_be_available_and_switch_to_it((By.CSS_SELECTOR, "iframe#cafe_main")))

        for current_page in range(start_page, end_page + 1):
            print(f"Scraping page {current_page}...")

            try:
                post_elements = driver.find_elements(By.XPATH, '//a[contains(@class, "article")]')
                if not post_elements:
                    print(f"No posts found on page {current_page}")
                    break

                for j in range(len(post_elements)):
                    post_elements = driver.find_elements(By.XPATH, '//a[contains(@class, "article")]')

                    try:
                        date_element_xpath = f'(//td[@class="td_date"])[{j+1}]'
                        date_element = WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.XPATH, date_element_xpath))
                        )
                        date_text = date_element.text
                        print(f"Post {j+1} has date: {date_text}")

                        # Parse the date string to a datetime object
                        post_date = datetime.strptime(date_text.strip(), "%Y.%m.%d.")
                        
                        # Check if the post date is from November 2024 onwards
                        if post_date < filter_date:
                            print(f"Skipping post {j+1} as it does not match the date filter (after November 2024)")
                            continue

                        # Extract the title directly from the anchor tag or child elements
                        post_element = post_elements[j]
                        title = post_element.text.strip()
                        print(f"Scraping post {j+1}: {title}")

                        # Click and proceed
                        post_element.click()
                        WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.XPATH, '//div[contains(@class, "se-main-container")]'))
                        )
                        time.sleep(2)

                        content_xpath = '//div[contains(@class, "se-main-container")]'
                        content_element = WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.XPATH, content_xpath))
                        )
                        content = content_element.text

                        comments, replies = extract_comments()

                        scraped_data.append({
                            "Date": date_text,
                            "Title": title,
                            "Content": content,
                            "Combined": f"{title} {content} " + " ".join(comments) + " " + " ".join(replies)
                        })

                        driver.back()
                        WebDriverWait(driver, 10).until(
                            EC.frame_to_be_available_and_switch_to_it((By.CSS_SELECTOR, "iframe#cafe_main"))
                        )

                    except (StaleElementReferenceException, TimeoutException, NoSuchElementException) as e:
                        print(f"Error with post {j+1} on page {current_page}: {e}")
                        continue

                next_page_link_xpath = f'//a[text()="{current_page + 1}"]'
                try:
                    next_page_element = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable((By.XPATH, next_page_link_xpath))
                    )
                    next_page_element.click()
                    time.sleep(2)
                except (TimeoutException, NoSuchElementException) as e:
                    print(f"No more pages or error clicking next page (page {current_page}): {e}")
                    break

            except (TimeoutException, NoSuchElementException) as e:
                print(f"Error with scraping on page {current_page}: {e}")
                break

    except TimeoutException as e:
        print(f"Error with initial search setup: {e}")

    return scraped_data

def extract_comments():
    comments = []
    replies = []

    try:
        comment_elements = driver.find_elements(By.XPATH, '//span[@class="text_comment"]')
        comments = [element.text for element in comment_elements]

        reply_elements = driver.find_elements(By.XPATH, '//li[contains(@class,"CommentItem--reply")]//span[@class="text_comment"]')
        replies = [element.text for element in reply_elements]

    except (TimeoutException, NoSuchElementException) as e:
        print(f"Error extracting comments or replies: {e}")

    return comments, replies

def analyze_with_gpt4(combined_texts):
    client = OpenAI(api_key=OpenAI.api_key)
    system_prompt = (
        "당신은 감정 분석과 요약에 능통한 전문가입니다. "
        "제공된 각 텍스트에 대해 감정 분류를 수행하세요 "
        "(긍정적, 중립적, 부정적) 그리고 주요 포인트와 의견을 요약하여 "
        "간단한 요약을 제공합니다."
    )
    
    summaries = []

    for text in combined_texts:
        chat_completion = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": text}
            ],
            max_tokens=4096
        )
        summary = chat_completion.choices[0].message.content.strip()
        summaries.append(summary)

    return summaries

def write_to_docx(summaries):
    doc = docx.Document()
    doc.add_heading('원더캠프 의견 요약', 0)

    for idx, summary in enumerate(summaries):
        doc.add_paragraph(f"Post {idx+1} Summary:", style='Heading2')
        doc.add_paragraph(summary)
    
    # Save in the current directory
    file_path = os.path.join(os.getcwd(), 'wonder_camp_summaries.docx')
    doc.save(file_path)
    print(f"Document saved to {file_path}")

def main():
    cafe_urls = [
        "https://cafe.naver.com/campingfirst",
        "https://cafe.naver.com/kzmkzmkzm/"
    ]
    search_keyword = "원더캠프"

    login_to_naver()  # Perform login once at the beginning

    all_data = []

    for idx, cafe_url in enumerate(cafe_urls):
        # Open each URL in a new tab
        if idx > 0:
            driver.execute_script("window.open('');")
            driver.switch_to.window(driver.window_handles[-1])

        # Navigate to the new URL in an existing session
        driver.get(cafe_url.strip())
        
        # Scrape data
        data = scrape_posts(search_keyword, start_page=1, end_page=3)
        all_data.extend(data)

    # Use GPT-4o to analyze sentiment and summarize
    combined_texts = [item["Combined"] for item in all_data]
    summaries = analyze_with_gpt4(combined_texts)

    # Write summaries to a docx file
    write_to_docx(summaries)
    print("Summaries have been generated and saved.")

if __name__ == "__main__":
    try:
        main()
    finally:
        driver.quit()