import os
import re
import pandas as pd
from openai import OpenAI
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.service import Service
from webdriver_manager.chrome import ChromeDriverManager
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException, StaleElementReferenceException
from dotenv import load_dotenv
from collections import defaultdict
from datetime import datetime
from docx import Document
import time

# Load environment variables from .env file
load_dotenv()

# Credentials
naver_marketing_id = os.getenv("NAVER_MARKETING_ID")
naver_marketing_password = os.getenv("NAVER_MARKETING_PASSWORD")
OpenAI.api_key = os.getenv("OPENAI_API_KEY")

# Set up WebDriver
options = Options()
options.add_argument("--disable-notifications")
options.add_argument('--no-sandbox')
service = Service(ChromeDriverManager().install())
driver = webdriver.Chrome(service=service, options=options)
driver.maximize_window()

def extract_product_info(excel_file):
    df = pd.read_excel(excel_file)
    product_category_mapping = {
        re.sub(r"[\[\(].*?[\]\)]", "", row['상품명']).strip(): row['카테고리']
        for _, row in df.iterrows()
    }
    return product_category_mapping

def login_to_naver():
    url = "https://nid.naver.com/nidlogin.login"
    driver.get(url)
    try:
        WebDriverWait(driver, 10).until(lambda d: d.find_element(By.ID, "id")).send_keys(naver_marketing_id)
        WebDriverWait(driver, 10).until(lambda d: d.find_element(By.ID, "pw")).send_keys(naver_marketing_password)
        WebDriverWait(driver, 10).until(lambda d: d.find_element(By.ID, "log.login")).click()
        time.sleep(20)  # Allow time for CAPTCHA manually
    except TimeoutException:
        print("Login elements did not load in time.")

def scrape_posts(search_keyword, start_page=1, end_page=3, processed_titles=set(), start_date=None, end_date=None):
    scraped_data = []

    try:
        search_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//input[@name="query"]'))
        )
        search_input.clear()
        search_input.send_keys(search_keyword)

        search_button = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//button[@onclick="searchBoard();return false;"]'))
        )
        search_button.click()
        time.sleep(2)

        WebDriverWait(driver, 10).until(EC.frame_to_be_available_and_switch_to_it((By.CSS_SELECTOR, "iframe#cafe_main")))

        if start_date:
            start_date = datetime.strptime(start_date, "%Y-%m-%d")
        if end_date:
            end_date = datetime.strptime(end_date, "%Y-%m-%d")

        current_page = start_page

        while current_page <= end_page:
            print(f"Scraping page {current_page}...")

            try:
                post_elements = driver.find_elements(By.XPATH, '//a[contains(@class, "article")]')
                if not post_elements:
                    print(f"No posts found on page {current_page}")
                    break

                for j in range(len(post_elements)):
                    post_elements = driver.find_elements(By.XPATH, '//a[contains(@class, "article")]')

                    try:
                        date_element_xpath = f'(//td[@class="td_date"])[{j+1}]'
                        date_element = WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.XPATH, date_element_xpath))
                        )
                        date_text = date_element.text
                        print(f"Post {j+1} has date: {date_text}")

                        if re.match(r"^\d{2}:\d{2}$", date_text):
                            today = datetime.today()
                            post_date = datetime.combine(today.date(), datetime.strptime(date_text, "%H:%M").time())
                        else:
                            post_date = datetime.strptime(date_text.strip(), "%Y.%m.%d.")

                        if (start_date and post_date < start_date) or (end_date and post_date > end_date):
                            print(f"Skipping post {j+1} as it does not match the date filter")
                            continue

                        post_element = post_elements[j]
                        title = post_element.text.strip()

                        # Skip posts already processed
                        if title in processed_titles:
                            print(f"Skipping post {j+1} with title '{title}' as it is already processed")
                            continue

                        print(f"Scraping post {j+1}: {title}")

                        post_element.click()
                        WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.XPATH, '//div[contains(@class, "se-main-container")]'))
                        )
                        time.sleep(2)

                        content_xpath = '//div[contains(@class, "se-main-container")]'
                        content_element = WebDriverWait(driver, 10).until(
                            EC.presence_of_element_located((By.XPATH, content_xpath))
                        )
                        content = content_element.text

                        comments, replies = extract_comments()

                        # Add post title to processed_titles to avoid duplicates
                        processed_titles.add(title)

                        scraped_data.append({
                            "Date": post_date.strftime("%Y-%m-%d %H:%M"),
                            "Title": title,
                            "Content": content,
                            "Combined": f"{title} {content} " + " ".join(comments) + " " + " ".join(replies)
                        })

                        driver.back()
                        WebDriverWait(driver, 10).until(
                            EC.frame_to_be_available_and_switch_to_it((By.CSS_SELECTOR, "iframe#cafe_main"))
                        )

                    except (StaleElementReferenceException, TimeoutException, NoSuchElementException) as e:
                        print(f"Error with post {j+1} on page {current_page}: {e}")
                        continue

                # Determine next page or next button actions
                if current_page % 10 == 0 and current_page < end_page:
                    next_button_xpath = '//a[@class="pgR"]/span[@class="m-tcol-c"][text()="다음"]'
                    try:
                        next_button_element = WebDriverWait(driver, 10).until(
                            EC.element_to_be_clickable((By.XPATH, next_button_xpath))
                        )
                        next_button_element.click()
                        time.sleep(2)
                    except (TimeoutException, NoSuchElementException) as e:
                        print(f"No more pages or error clicking next button (page {current_page}): {e}")
                        break
                elif current_page < end_page:
                    next_page_link_xpath = f'//a[text()="{current_page + 1}"]'
                    try:
                        next_page_element = WebDriverWait(driver, 10).until(
                            EC.element_to_be_clickable((By.XPATH, next_page_link_xpath))
                        )
                        next_page_element.click()
                        time.sleep(2)
                    except (TimeoutException, NoSuchElementException) as e:
                        print(f"No more pages or error clicking next page (page {current_page}): {e}")
                        break
                current_page += 1

            except (TimeoutException, NoSuchElementException) as e:
                print(f"Error with scraping on page {current_page}: {e}")
                break

    except TimeoutException as e:
        print(f"Error with initial search setup: {e}")

    return scraped_data

def extract_comments():
    comments = []
    replies = []

    try:
        comment_elements = driver.find_elements(By.XPATH, '//span[@class="text_comment"]')
        comments = [element.text for element in comment_elements]

        reply_elements = driver.find_elements(By.XPATH, '//li[contains(@class,"CommentItem--reply")]//span[@class="text_comment"]')
        replies = [element.text for element in reply_elements]

    except (TimeoutException, NoSuchElementException) as e:
        print(f"Error extracting comments or replies: {e}")

    return comments, replies

def analyze_with_gpt4o(all_data, product_category_mapping):
    client = OpenAI(api_key=OpenAI.api_key)
    system_prompt = (
        """당신은 감정 분석과 마케팅에 유능한 전문가입니다. 주어진 각 텍스트에 대해 다음을 수행하세요:
    
        1. 주어진 목록에서 각 제품명이 몇 번 언급되었는지 식별합니다.
        2. 각 제품이 속한 카테고리가 몇 번 언급되었는지 기록합니다.
        3. 포스트, 댓글 및 대댓글의 전반적인 감정을 긍정적, 부정적, 중립적으로 분석하고 분류합니다.
    
        응답은 감정 분류를 구조화된 형식으로 제공하세요 (예: 감정: 긍정적/부정적/중립적)."""
    )

    product_count = defaultdict(int)
    category_count = defaultdict(int)
    sentiment_count = {'positive': 0, 'neutral': 0, 'negative': 0}
    total_posts = len(all_data)

    summaries = []

    for item in all_data:
        text = item["Combined"]
        try:
            chat_completion = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": text}
                ],
                max_tokens=8192
            )
            summary = chat_completion.choices[0].message.content.strip()
            summaries.append(summary)

            # Count product mentions per category: each product is counted only once per post
            mentioned_products = set()
            for product, category in product_category_mapping.items():
                product_regex = re.compile(re.escape(product.replace(" ", "")), re.IGNORECASE)
                if product_regex.search(text.replace(" ", "")) and product not in mentioned_products:
                    mentioned_products.add(product)
                    product_count[product] += 1
                    category_count[category] += 1

            # Extract sentiment from the structured response
            if "감정: 긍정적" in summary:
                sentiment_count['positive'] += 1
            elif "감정: 부정적" in summary:
                sentiment_count['negative'] += 1
            elif "감정: 중립적" in summary:
                sentiment_count['neutral'] += 1
            else:
                print(f"Sentiment not detected correctly in summary: {summary}")

        except Exception as e:
            print(f"Error during GPT-4o call: {e}")
            summaries.append("Error processing text")

    return summaries, dict(product_count), dict(category_count), dict(sentiment_count), total_posts

def write_summary_to_docx(product_count, category_count, sentiment_count, total_posts, all_data, summaries, output_file):
    """
    Writes a summarized analysis of the data and sentiment to a DOCX file.
    """
    doc = Document()

    # Add a title for the document
    doc.add_heading('Sentiment Analysis Summary', level=1)

    # Create table for product mentions (sorted alphabetically by product name)
    doc.add_heading('Product Mentions', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product'
    hdr_cells[1].text = 'Count'
    
    for product in sorted(product_count):
        row_cells = table.add_row().cells
        row_cells[0].text = product
        row_cells[1].text = str(product_count[product])

    # Create table for category mentions (sorted alphabetically by category name)
    doc.add_heading('Category Mentions', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Total Mentions'
    
    for category in sorted(category_count):
        row_cells = table.add_row().cells
        row_cells[0].text = category
        row_cells[1].text = str(category_count[category])

    # Create table for sentiment analysis
    doc.add_heading('Sentiment Analysis', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Sentiment'
    hdr_cells[1].text = 'Count'
    for sentiment, count in sentiment_count.items():
        row_cells = table.add_row().cells
        row_cells[0].text = sentiment
        row_cells[1].text = str(count)

    # Total posts overview
    doc.add_heading('Overview', level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    row_cells = table.add_row().cells
    row_cells[0].text = 'Total Posts'
    row_cells[1].text = str(total_posts)

    # Add detailed summaries with titles at the end
    doc.add_heading('Detailed Summaries', level=2)
    for item, summary in zip(all_data, summaries):
        # Add the post title as a heading
        doc.add_heading(f"Title: {item['Title']}", level=3)
        doc.add_paragraph(summary)
        doc.add_paragraph("\n")  # Space between summaries

    # Save the document
    doc.save(output_file)
    print(f"Results saved to {output_file}")

def main():
    cafe_urls = [
        "https://cafe.naver.com/campingfirst",
    ]
    
    # Example with multiple keywords - ensure they are split correctly
    search_keywords = "'카즈미', 'KZM', 'KZM OUTDOOR'"
    search_keywords_list = [kw.strip().strip("'") for kw in search_keywords.split(',')]

    login_to_naver()

    all_data = []
    processed_titles = set()  # Set to keep track of processed post titles

    start_date = "2024-11-01"  # Example start date
    end_date = "2024-11-30"    # Example end date

    for keyword in search_keywords_list:
        for idx, cafe_url in enumerate(cafe_urls):
            print(f"Searching for keyword: {keyword} in cafe: {cafe_url}")
            if idx > 0:
                driver.execute_script("window.open('');")
                driver.switch_to.window(driver.window_handles[-1])

            driver.get(cafe_url.strip())
            
            # Scrape posts for each keyword individually
            data = scrape_posts(
                search_keyword=keyword,
                start_page=1,
                end_page=22,
                processed_titles=processed_titles,
                start_date=start_date,
                end_date=end_date
            )
            all_data.extend(data)

    # Load product category mapping from Excel
    excel_file_path = r"C:\Users\jung\Desktop\AICC\playauto\product\items.xls"
    product_category_mapping = extract_product_info(excel_file_path)
    
    # Analyze scraped data with GPT-4o
    summaries, product_count, category_count, sentiment_count, total_posts = analyze_with_gpt4o(
        all_data, product_category_mapping
    )

    # Write summarized analysis to a DOCX file
    write_summary_to_docx(
        product_count, category_count, sentiment_count, total_posts, all_data, summaries, 'summary_results.docx'
    )
    print("Summaries have been generated and saved.")

if __name__ == "__main__":
    try:
        main()
    finally:
        driver.quit()